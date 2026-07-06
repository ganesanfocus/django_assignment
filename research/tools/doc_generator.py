# tools/doc_generator.py
import os
from fpdf import FPDF
from docx import Document

def create_output_dir():
    if not os.path.exists("output"):
        os.makedirs("output")

def generate_pdf(content: str, filename: str = "output/report.pdf"):
    create_output_dir()
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    # Replace unsupported characters for basic latin-1 FPDF
    safe_content = content.encode('latin-1', 'replace').decode('latin-1')
    
    pdf.multi_cell(0, 10, txt=safe_content)
    pdf.output(filename)
    return filename

def generate_docx(content: str, filename: str = "output/report.docx"):
    create_output_dir()
    doc = Document()
    doc.add_heading('Research Report', 0)
    doc.add_paragraph(content)
    doc.save(filename)
    return filename

def generate_document(content: str, format_type: str = "pdf") -> str:
    """Generates a document based on the requested format."""
    format_type = format_type.lower()
    if "docx" in format_type or "word" in format_type:
        path = generate_docx(content)
        return f"Document successfully saved as DOCX at: {path}"
    else:
        path = generate_pdf(content)
        return f"Document successfully saved as PDF at: {path}"